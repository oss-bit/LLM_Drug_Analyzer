from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader, WebBaseLoader, AsyncHtmlLoader
from langchain.vectorstores import Chroma
from langchain.chains.qa_with_sources.retrieval import RetrievalQAWithSourcesChain
from langchain.schema import Document
import chainlit as cl
from langchain_groq import ChatGroq
from langchain.embeddings.base import Embeddings
from sentence_transformers import SentenceTransformer
from typing import List
import os
import docx


class SentenceTransformerEmbeddings(Embeddings):
    def __init__(self, model_name: str):
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, documents: List[str]) -> List[List[float]]:
        documents = [str(doc) for doc in documents]
        return self.model.encode(documents).tolist()

    def embed_query(self, query: str) -> List[float]:
        if not isinstance(query, str):
            query = str(query)
        return self.model.encode([query])[0].tolist()
    


embedding = SentenceTransformerEmbeddings("sentence-transformers/all-MiniLM-L6-v2")

 # Add your Groq API key here
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
llm = ChatGroq(temperature=0, model_name='llama3-8b-8192')

# System prompt for research topic recommendation
SYSTEM_PROMPT = """You are an expert pharmaceutical advisor and drug information specialist with deep knowledge in pharmacology, clinical medicine, and drug safety. You help people understand medications by analyzing drug literature, clinical data, and pharmaceutical resources.

Given retrieved content from drug databases, medical literature, and pharmaceutical sources, your responsibilities are:

• Analyze drug mechanisms, pharmacokinetics, and pharmacodynamics
• Identify drug interactions, contraindications, and safety profiles  
• Explain therapeutic uses, dosing guidelines, and clinical applications
• Highlight side effects, adverse reactions, and monitoring requirements
• Compare different medications and treatment options when relevant
• Provide evidence-based recommendations with proper medical context

Context from drug databases and medical literature:
{summaries}

Question: {question}

When the provided content lacks specific information about a medication, do not say "I don't know" or mention that the content is insufficient. Instead, provide comprehensive information based on general pharmacological knowledge, established drug classes, or related therapeutic areas. Always offer helpful, accurate medical information, even if you must draw from broader pharmaceutical expertise. Make it clear when you are providing general drug class information versus specific product details.

Structure your response as:

**Drug Overview:**
- Generic/Brand names and drug classification
- Primary mechanism of action and how it works in the body

**Clinical Applications:**
- Approved uses and therapeutic indications
- Typical dosing and administration guidelines

**Safety Profile:**
- Common and serious side effects
- Contraindications and precautions
- Drug interactions and monitoring requirements

**Clinical Considerations:**
- Special populations (elderly, pregnancy, pediatric)
- Comparative effectiveness with alternative treatments
- Important counseling points for patients

**Sources Referenced:**
- Cite specific sources that informed your analysis

Always include appropriate medical disclaimers and encourage consultation with healthcare professionals for personalized medical advice.

Answer:"""

# Predefined sources - Add your sources here
PREDEFINED_SOURCES = {
    # 'pdfs': [
    #     # Add paths to your PDF files here
    #     # 'path/to/research_paper1.pdf',
    #     # 'path/to/research_paper2.pdf',
    # ],
    # 'word_docs': [
    #     # Add paths to your Word documents here
    #     # 'path/to/document1.docx',
        # 'path/to/document2.doc',
    # ],
    'web_urls': [
        'https://en.wikipedia.org/wiki/Ibuprofen',
        'https://en.wikipedia.org/wiki/Paracetamol',  # Fixed: removed leading space
        'https://en.wikipedia.org/wiki/Aspirin',
        'https://en.wikipedia.org/wiki/Diclofenac',
        'https://en.wikipedia.org/wiki/Tramadol',
        'https://en.wikipedia.org/wiki/Naproxen'
    ]
}


# LangChain provides excellent web loaders - no need for custom implementation


def load_pdf_documents(pdf_paths: List[str]) -> List[Document]:
    """Load PDF documents from file paths"""
    documents = []
    for pdf_path in pdf_paths:
        if os.path.exists(pdf_path):
            try:
                loader = PyPDFLoader(pdf_path)
                docs = loader.load()
                for i, doc in enumerate(docs):
                    doc.metadata['source'] = f"{os.path.basename(pdf_path)}_page_{i+1}"
                    doc.metadata['type'] = 'pdf'
                documents.extend(docs)
            except Exception as e:
                print(f"Error loading PDF {pdf_path}: {str(e)}")
    return documents


def load_word_documents(word_paths: List[str]) -> List[Document]:
    """Load Word documents from file paths"""
    documents = []
    for word_path in word_paths:
        if os.path.exists(word_path):
            try:
                # Try using python-docx first for .docx files
                if word_path.endswith('.docx'):
                    doc_content = docx.Document(word_path)
                    text = '\n'.join([paragraph.text for paragraph in doc_content.paragraphs])
                    doc = Document(
                        page_content=text,
                        metadata={'source': os.path.basename(word_path), 'type': 'word'}
                    )
                    documents.append(doc)
                else:
                    # Use UnstructuredWordDocumentLoader for .doc files
                    loader = UnstructuredWordDocumentLoader(word_path)
                    docs = loader.load()
                    for doc in docs:
                        doc.metadata['source'] = os.path.basename(word_path)
                        doc.metadata['type'] = 'word'
                    documents.extend(docs)
            except Exception as e:
                print(f"Error loading Word document {word_path}: {str(e)}")
    return documents


def load_web_documents(urls: List[str]) -> List[Document]:
    """Load documents from web URLs using LangChain's WebBaseLoader"""
    documents = []
    
    if not urls:
        return documents
    
    try:
        # Use LangChain's WebBaseLoader for better web scraping
        loader = WebBaseLoader(urls)
        
        # Configure loader settings for better content extraction
        loader.requests_kwargs = {
            'timeout': 30,
            'headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
        }
        
        docs = loader.load()
        
        # Add proper metadata
        for doc in docs:
            doc.metadata['type'] = 'web'
            if 'source' not in doc.metadata:
                doc.metadata['source'] = doc.metadata.get('url', 'unknown_url')
        
        documents.extend(docs)
        print(f"Successfully loaded {len(docs)} web documents")
        
    except Exception as e:
        print(f"Error loading web documents: {str(e)}")
        
        # Fallback: try loading each URL individually
        for url in urls:
            try:
                loader = WebBaseLoader([url])
                loader.requests_kwargs = {
                    'timeout': 10,
                    'headers': {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    }
                }
                docs = loader.load()
                for doc in docs:
                    doc.metadata['type'] = 'web'
                    doc.metadata['source'] = url
                documents.extend(docs)
                print(f"Loaded web document from: {url}")
            except Exception as individual_error:
                print(f"Failed to load {url}: {str(individual_error)}")
    
    return documents


def initialize_knowledge_base():
    """Initialize the knowledge base with predefined sources"""
    all_documents = []
    
    # # Load PDF documents
    # if PREDEFINED_SOURCES['pdfs']:
    #     pdf_docs = load_pdf_documents(PREDEFINED_SOURCES['pdfs'])
    #     all_documents.extend(pdf_docs)
    #     print(f"Loaded {len(pdf_docs)} PDF documents")
    
    # # Load Word documents
    # if PREDEFINED_SOURCES['word_docs']:
    #     word_docs = load_word_documents(PREDEFINED_SOURCES['word_docs'])
    #     all_documents.extend(word_docs)
    #     print(f"Loaded {len(word_docs)} Word documents")
    
    # Load web documents
    if PREDEFINED_SOURCES['web_urls']:
        web_docs = load_web_documents(PREDEFINED_SOURCES['web_urls'])
        all_documents.extend(web_docs)
        print(f"Loaded {len(web_docs)} web documents")
    
    if not all_documents:
        print("No documents loaded. Please add sources to PREDEFINED_SOURCES.")
        return None, None
    
    # Split documents
    # splits = text_splitter.split_documents(all_documents)
    splits = text_splitter.split_documents([doc if isinstance(doc, Document) else Document(page_content=str(doc)) for doc in all_documents])

    
    # Ensure each split has proper metadata
    for i, split in enumerate(splits):
        if 'source' not in split.metadata:
            split.metadata['source'] = f'document_chunk_{i}'
    
    # Create vector store
    texts = [doc.page_content for doc in splits]
    metadatas = [doc.metadata for doc in splits]
    
    vec_search = Chroma.from_texts(
        texts=texts,
        embedding=embedding,
        metadatas=metadatas,
        collection_name="research_recommender_store"
    )
    
    return vec_search, splits


def create_custom_prompt(user_query: str, additional_prompts: List[str] = None) -> str:
    """Create a custom prompt combining system prompt, additional prompts, and user query"""
    full_prompt = SYSTEM_PROMPT
    
    if additional_prompts:
        full_prompt += "\n\nAdditional Instructions:\n"
        for prompt in additional_prompts:
            full_prompt += f"- {prompt}\n"
    
    full_prompt += f"\n\nUser Query: {user_query}"
    return full_prompt


@cl.on_chat_start
async def start():
    await cl.Message(content="Drug Analyzer Initializing...").send()
    
    # Initialize knowledge base
    vec_search, docs = await cl.make_async(initialize_knowledge_base)()
    
    if vec_search is None:
        await cl.Message(content="❌ Failed to initialize knowledge base. Please add sources to PREDEFINED_SOURCES in the code.").send()
        return
    
    # Store in session
    cl.user_session.set("vec_search", vec_search)
    cl.user_session.set("docs", docs)
    
    # Create QA chain
    chain = RetrievalQAWithSourcesChain.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vec_search.as_retriever(search_kwargs={"k": 8})
    )
    cl.user_session.set("chain", chain)
    
    await cl.Message(content=f"Drug Analyzer Ready!").send()


@cl.on_message
async def main(message):
    chain = cl.user_session.get('chain')
    docs = cl.user_session.get("docs")
    
    if not chain:
        await cl.Message(content="Please wait for the system to initialize completely.").send()
        return
    
    # You can add additional system prompts here
    additional_prompts = [
        "Focus on topics that have practical applications and societal impact",
        "Consider interdisciplinary research opportunities",
        "Highlight any methodological innovations that could be explored"
    ]
    
    # Create custom prompt
    custom_query = create_custom_prompt(message.content, additional_prompts)
    
    # Set up callback for streaming
    cb = cl.AsyncLangchainCallbackHandler(
        stream_final_answer=True,
        answer_prefix_tokens=["FINAL", "ANSWER"]
    )
    cb.answer_reached = True
    
    # Get response from chain
    res = await chain.acall(custom_query, callbacks=[cb])
    answer = res['answer']
    sources = res['sources'].strip()
    
    # Process sources
    source_elements = []
    docs_metadata = [doc.metadata for doc in docs]
    all_sources = [m['source'] for m in docs_metadata]
    
    if sources:
        found_sources = []
        for source in sources.split(','):
            source_name = source.strip().replace('.', '')
            try:
                index = all_sources.index(source_name)
            except ValueError:
                continue
            text = docs[index].page_content
            found_sources.append(source_name)
            source_elements.append(cl.Text(content=text, name=source_name))
        
        if found_sources:
            answer += f'\n\n📚 **Sources Used:**\n{", ".join(found_sources)}'
        else:
            answer += '\n\n📚 **Sources Used:** No specific sources identified'
    
    # Send response
    if cb.has_streamed_final_answer:
        cb.final_stream.elements = source_elements
        await cb.final_stream.update()
    else:
        await cl.Message(content=answer, elements=source_elements).send()
